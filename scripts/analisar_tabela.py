"""
Script para analisar a estrutura da tabela do template
"""
import sys
import os
from docx import Document

if os.name == 'nt':
    sys.stdout.reconfigure(encoding='utf-8')

doc = Document("OFERTA-DO-DIA.docx")

print(f"Total de tabelas: {len(doc.tables)}")

if doc.tables:
    tabela = doc.tables[0]
    print(f"Linhas: {len(tabela.rows)}")
    print(f"Colunas: {len(tabela.columns)}")
    
    # Mostrar conteúdo das primeiras 5 linhas para entender o cabeçalho
    for i, row in enumerate(tabela.rows[:5]):
        texto_celulas = [cell.text.strip() for cell in row.cells]
        print(f"Linha {i}: {texto_celulas}")
