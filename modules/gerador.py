"""
Módulo Gerador - Gera documento OFERTA-DO-DIA.docx
"""
from docx import Document
from pathlib import Path
from datetime import datetime
import os

class GeradorOferta:
    """Classe para gerar documento OFERTA-DO-DIA"""
    
    def __init__(self, produtos):
        self.produtos = produtos
        self.template_path = None
        self.output_dir = "output"
    
    def gerar_docx(self, template_path="OFERTA-DO-DIA.docx", output_path="output/OFERTA-DO-DIA.docx"):
        """
        Gera documento DOCX com produtos filtrados
        
        Args:
            template_path (str): Caminho do template DOCX
            output_path (str): Caminho de saída do DOCX gerado
            
        Returns:
            str: Caminho do arquivo gerado
        """
        self.template_path = template_path
        self.output_dir = os.path.dirname(output_path)
        
        try:
            # Criar novo documento ou usar template existente
            if Path(template_path).exists():
                doc = Document(template_path)
                print(f"[OK] Template carregado: {template_path}")
            else:
                doc = Document()
                self._criar_template_basico(doc)
                print("[INFO] Criando documento novo (template nao encontrado)")
                self.template_path = None # Marca que não usou template
            
            # Adicionar produtos
            self._adicionar_produtos(doc)
            
            # Garantir que diretório existe
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            
            # Salvar documento
            doc.save(output_path)
            print(f"[OK] DOCX gerado: {output_path}")
            
            # Atualizar data de validade via XML (após salvar)
            # Isso é feito DEPOIS de salvar porque precisamos editar o arquivo no disco
            if self.template_path:
                self._adicionar_data_validade(output_path)
            
            return output_path
            
        except Exception as e:
            print(f"[ERRO] Erro ao gerar DOCX: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _adicionar_data_validade(self, output_path):
        """
        Atualiza a data de validade editando diretamente o XML do DOCX
        Isso é necessário porque o python-docx as vezes não encontra textos em headers complexos
        """
        import zipfile
        import re
        import tempfile
        import shutil
        
        print("[INFO] Tentando atualizar data via XML direto...")
        
        # Criar diretório temporário
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Extrair DOCX
            with zipfile.ZipFile(output_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Data de hoje
            data_hoje = datetime.now().strftime("%d/%m/%Y")
            data_hoje_curta = datetime.now().strftime("%d/%m/%y")
            
            # Padrão para encontrar datas: DD/MM/YY ou DD/MM/YYYY
            # Procura datas que estejam no mesmo arquivo XML que "VALIDO"
            padrao_data = r'(\d{2}/\d{2}/\d{2,4})'
            
            # Arquivos para verificar
            arquivos_modificados = 0
            
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    if file.endswith('.xml'):
                        file_path = os.path.join(root, file)
                        
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                        
                        # Verificar se tem "VALIDO" neste arquivo
                        if 'VALIDO' in content.upper() or 'VÁLIDO' in content.upper():
                            # Procurar datas neste arquivo
                            datas_encontradas = re.findall(padrao_data, content)
                            
                            if datas_encontradas:
                                novo_content = content
                                for data_antiga in datas_encontradas:
                                    # Evitar substituir a própria data de hoje se já estiver certa
                                    if data_antiga == data_hoje or data_antiga == data_hoje_curta:
                                        continue
                                        
                                    # Substituir pela data de hoje (mantendo formato curto/longo)
                                    if len(data_antiga) == 8: # DD/MM/YY
                                        novo_content = novo_content.replace(data_antiga, data_hoje_curta)
                                        print(f"[XML] Substituindo {data_antiga} por {data_hoje_curta} em {file}")
                                    else: # DD/MM/YYYY
                                        novo_content = novo_content.replace(data_antiga, data_hoje)
                                        print(f"[XML] Substituindo {data_antiga} por {data_hoje} em {file}")
                                
                                if novo_content != content:
                                    with open(file_path, 'w', encoding='utf-8') as f:
                                        f.write(novo_content)
                                    arquivos_modificados += 1
            
            if arquivos_modificados > 0:
                # Recompactar DOCX
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                    for root, dirs, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, temp_dir)
                            zip_out.write(file_path, arcname)
                print(f"[OK] Data atualizada com sucesso via XML em {arquivos_modificados} arquivos!")
            else:
                print("[INFO] Nenhuma data antiga encontrada para substituir no XML")
                
        except Exception as e:
            print(f"[ERRO] Falha ao editar XML: {e}")
        finally:
            # Limpar temporários
            shutil.rmtree(temp_dir)

    
    def _criar_template_basico(self, doc):
        """Cria template básico se não existir"""
        # Título
        titulo = doc.add_heading('OFERTA DO DIA', 0)
        titulo.alignment = 1  # Centralizado
        
        # Data e Validade
        data_hoje = datetime.now().strftime("%d/%m/%Y")
        p_data = doc.add_paragraph(f"Data: {data_hoje}")
        p_data.alignment = 1
        
        p_validade = doc.add_paragraph(f"Válido para: {data_hoje}")
        p_validade.alignment = 1
        
        # Formatação em negrito para validade
        for run in p_validade.runs:
            run.bold = True
        
        doc.add_paragraph("")  # Espaço
    
    def _adicionar_produtos(self, doc):
        """Adiciona produtos à tabela existente no template"""
        
        # Procurar tabela existente
        if not doc.tables:
            print("[ERRO] Nenhuma tabela encontrada no template!")
            return

        tabela = doc.tables[0] # Usa a primeira tabela
        print(f"[INFO] Tabela encontrada: {len(tabela.rows)} linhas, {len(tabela.columns)} colunas")
        
        # Identificar linha de início dos dados
        # Assume que as 2 primeiras linhas são cabeçalho (baseado na análise: 'HIGIENE...' e 'NOME...')
        linha_inicio = 2 
        
        # Limpar dados existentes (da linha_inicio em diante)
        # Se houver muitas linhas vazias no template, vamos usá-las
        # Se houver dados antigos, vamos sobrescrever
        
        total_produtos = len(self.produtos)
        total_linhas_tabela = len(tabela.rows)
        
        for i, produto in enumerate(self.produtos):
            indice_linha = linha_inicio + i
            
            # Se a linha já existe, usa ela
            if indice_linha < total_linhas_tabela:
                linha = tabela.rows[indice_linha]
            else:
                # Se não existe, cria nova
                linha = tabela.add_row()
            
            # Preencher células (Mapeamento: 0=Desc, 1=Unid, 2=Preço)
            celulas = linha.cells
            
            # Garantir 3 colunas
            if len(celulas) >= 3:
                # Descrição
                celulas[0].text = str(produto['descricao'])
                # Unidade (limpar 'CX' duplicado se vier da extração, ex: '255 CX')
                # O template pede UNIDADE (ex: CX, FD, UN). 
                # A extração retorna 'estoque' (ex: 955) e 'unidade' (ex: CX).
                # O usuário quer apenas a UNIDADE na coluna 2? Ou Estoque?
                # Na imagem 1: Coluna 2 tem "CX". Coluna 3 tem "R$ 118.00".
                # O estoque NÃO aparece na imagem.
                celulas[1].text = str(produto['unidade'])
                
                # Preço
                celulas[2].text = f"R$ {produto['preco']:.2f}"
                
                # Centralizar Unidade e Preço (opcional, mas fica bonito)
                for p in celulas[1].paragraphs: p.alignment = 1 # Center
                for p in celulas[2].paragraphs: p.alignment = 1 # Center
        
        # Limpar linhas excedentes (se o template tiver mais linhas que produtos)
        ultima_linha_preenchida = linha_inicio + total_produtos
        
        # Não podemos deletar linhas facilmente no python-docx sem mexer no XML
        # Então vamos apenas limpar o texto das linhas sobrando para ficarem vazias
        for i in range(ultima_linha_preenchida, total_linhas_tabela):
            linha = tabela.rows[i]
            for celula in linha.cells:
                celula.text = ""
        
        print(f"[OK] {total_produtos} produtos inseridos na tabela (Colunas: Descrição, Unidade, Preço)")

if __name__ == "__main__":
    # Teste com dados de exemplo
    produtos_exemplo = [
        {
            'codigo': '9576',
            'descricao': 'TESTE XML UPDATE',
            'estoque': 999,
            'unidade': 'CX',
            'preco': 100.00
        }
    ]
    
    gerador = GeradorOferta(produtos_exemplo)
    gerador.gerar_docx()
